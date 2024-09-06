from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.models import User
from django.utils import timezone
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.forms import UserCreationForm, AuthenticationForm
from django.urls import reverse
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from . import models
from . import forms
from random import choice, sample
from colorama import Fore, init
from io import BytesIO
import openpyxl
from docx import Document
import warnings
import datetime

#def index(request):
    #return render(request, 'index.html')

init(autoreset=True)

@login_required(login_url='login')
def quizList(request):
    images = [
        'https://st2.depositphotos.com/2769299/7314/i/450/depositphotos_73146775-stock-photo-a-stack-of-books-on.jpg',
        'https://img.freepik.com/free-photo/creative-composition-world-book-day_23-2148883765.jpg',
        'https://profit.pakistantoday.com.pk/wp-content/uploads/2018/04/Stack-of-books-great-education.jpg',
        'https://live-production.wcms.abc-cdn.net.au/73419a11ea13b52c6bd9c0a69c10964e?impolicy=wcms_crop_resize&cropH=1080&cropW=1918&xPos=1&yPos=0&width=862&height=485',
        'https://live-production.wcms.abc-cdn.net.au/398836216839841241467590824c5cf1?impolicy=wcms_crop_resize&cropH=2813&cropW=5000&xPos=0&yPos=0&width=862&height=485',
        'https://images.theconversation.com/files/45159/original/rptgtpxd-1396254731.jpg?ixlib=rb-4.1.0&q=45&auto=format&w=1356&h=668&fit=crop'
    ]
    
    quizes = models.Quiz.objects.filter(author=request.user)
    # images = sample(len(quizes), images)

    quizes_list = []

    for quiz in quizes:
        quiz.img = choice(images)
        quizes_list.append(quiz)

    return render(request, 'quiz-list.html', {'quizes':quizes_list})

@login_required(login_url='login')
def quizDetail(request, slug):
    quiz = get_object_or_404(models.Quiz, slug=slug)
    return render(request, 'quiz-detail.html', {'quiz': quiz})
    
    
@login_required(login_url='login')
def take_quiz(request, slug):
    quiz = get_object_or_404(models.Quiz, slug=slug)
    
    if request.method == 'POST':
        form = forms.QuizForm(request.POST, quiz=quiz)
        if form.is_valid():
            answer = models.Answer.objects.create(
                quiz=quiz,
                author=request.user,
                start_time = datetime.datetime.now() 
            )
            
            for question in quiz.questions.all():
                option_id = form.cleaned_data.get(f'question_{question.id}')
                option = models.Option.objects.get(id=option_id)
                models.AnswerDetail.objects.create(
                    answer=answer,
                    question=question,
                    user_choice=option
                )
            
            answer.end_time = datetime.datetime.now()
            answer.save()
            
            return redirect('quiz-results', slug=quiz.slug)

    else:
        form = forms.QuizForm(quiz=quiz)
    
    return render(request, 'quiz_form.html', {'form': form, 'quiz': quiz})
    
    
@login_required(login_url='login')
def quiz_results(request, slug):
    quiz = get_object_or_404(models.Quiz, slug=slug)
    results = models.Answer.objects.filter(quiz=quiz, author=request.user)
    
    print(Fore.RED + str(f"quiz: {quiz}"))
    for i in results:
        print(Fore.YELLOW + str(f"results: {i.correct_percentage}"))
    
    return render(request, 'quiz_results.html', {'quiz': quiz, 'results': results})


@login_required(login_url='login')
def questionDelete(request, id, slug):
    models.Question.objects.get(id=id).delete()
    return redirect('quizDetail', slug=slug)


@login_required(login_url='login')
def createQuiz(request):
    if request.method == 'POST':
        quiz = models.Quiz.objects.create(
            name=request.POST['name'],
            amount=request.POST['amount'],
            author=request.user
        )
        return redirect('quizDetail', slug=quiz.slug)
    return render(request, 'quiz-create.html')


@login_required(login_url='login')
def questionCreate(request, slug):
    quiz = get_object_or_404(models.Quiz, slug=slug)
    if request.method == 'POST':
        question_text = request.POST['name']
        true = request.POST['true']
        false_list = request.POST.getlist('false-list')

        question = models.Question.objects.create(
            name=question_text,
            quiz=quiz,
        )
        models.Option.objects.create(
            question=question,
            name=true,
            correct=True,
        )

        for false in false_list:
            models.Option.objects.create(
                name=false,
                question=question,
            )

        return redirect('quizDetail', slug=quiz.slug)

    return render(request, 'question-create.html', {'quiz': quiz})


@login_required(login_url='login')
def questionDetail(request, id):
    question = models.Question.objects.get(id=id)
    return render(request, 'question-detail.html', {'question':question})


@login_required(login_url='login')
def deleteOption(request, ques, option):
    question = models.Question.objects.get(id=ques)
    models.Option.objects.get(question=question, id=option).delete()
    return redirect('questionDetail', id=ques)



@login_required(login_url='login')
def detail(request):
    quizzes = models.Quiz.objects.all()
    results = []

    for quiz in quizzes:
        total_questions = quiz.questions_count

        answer_details = models.AnswerDetail.objects.filter(answer__quiz=quiz)

        correct_answers = answer_details.filter(
            user_choice__correct=True
        ).count()
        incorrect_answers = answer_details.filter(
            user_choice__correct=False
        ).count()

        results.append({
            'quiz': quiz,
            'total_questions': total_questions,
            'correct_answers': correct_answers,
            'incorrect_answers': incorrect_answers
        })
        
    print(Fore.RED + str(f"{results}"))

    return render(request, 'detail.html', {'results': results})
    
   
@login_required(login_url='login')
def results_detail(request, slug):
    quiz = get_object_or_404(models.Quiz, slug=slug)
    answers = models.Answer.objects.filter(quiz=quiz)

    details = []
    for answer in answers:
        answer_details = models.AnswerDetail.objects.filter(answer=answer)
        correct_answers = answer_details.filter(is_correct=True).count()
        incorrect_answers = answer_details.filter(is_correct=False).count()

        total_questions = correct_answers + incorrect_answers
        correct_percentage = (correct_answers / total_questions) * 100 if total_questions > 0 else 0

        details.append({
            'user': answer.author,
            'correct_answers': correct_answers,
            'incorrect_answers': incorrect_answers,
            'correct_percentage': correct_percentage
        })

    return render(request, 'results_detail.html', {'quiz': quiz, 'details': details})

    

@login_required(login_url='login')  
def participant_results(request, quiz_id):
    quiz = get_object_or_404(models.Quiz, id=quiz_id)
    participants = models.Answer.objects.filter(quiz=quiz).values('author').distinct()

    participant_details = []

    for participant in participants:
        user = User.objects.get(id=participant['author'])
        
        total_answers = models.AnswerDetail.objects.filter(
            answer__author=user,
            answer__quiz=quiz
        ).count()
        
        correct_answers = models.AnswerDetail.objects.filter(
            answer__author=user,
            answer__quiz=quiz,
            user_choice__correct=True
        ).count()
        
        incorrect_answers = models.AnswerDetail.objects.filter(
            answer__author=user,
            answer__quiz=quiz,
            user_choice__correct=False
        ).count()
        
        correct_percentage = (correct_answers / total_answers) * 100 if total_answers > 0 else 0
        
        participant_details.append({
            'user': user,
            'correct_answers': correct_answers,
            'incorrect_answers': incorrect_answers,
            'correct_percentage': correct_percentage
        })

    return render(request, 'participant_results.html', {'quiz': quiz, 'participant_details': participant_details})
    
    
@login_required(login_url='login')  
def participant_detail(request, quiz_id, user_id):
    quiz = get_object_or_404(models.Quiz, id=quiz_id)
    user = get_object_or_404(User, id=user_id)
    details = models.AnswerDetail.objects.filter(answer__author=user, answer__quiz=quiz)
    
    return render(request, 'participant_detail.html', {
        'quiz': quiz,
        'user': user,
        'details': details
    })
    
    
@login_required(login_url='login')
def export_results_to_excel(request, slug):
    quiz = get_object_or_404(models.Quiz, slug=slug)
    results = models.Answer.objects.filter(quiz=quiz, author=request.user)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Quiz Results"
    
    ws.append(['#', 'Name', 'Correct Answers', 'Incorrect Answers', 'Correct Percentage'])
    
    for result in results:
        ws.append([
            result.id,
            result.author.username,
            result.correct_answers,
            result.incorrect_answers,
            f"{result.correct_percentage:.2f}%"
        ])
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{quiz.name}_results.xlsx"'
    wb.save(response)
    return response
    
    
@login_required(login_url='login')
def export_results_to_word(request, slug):
    quiz = get_object_or_404(models.Quiz, slug=slug)
    results = models.Answer.objects.filter(quiz=quiz, author=request.user)
    
    doc = Document()
    doc.add_heading(f'{quiz.name} Results', level=1)
    
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '#'
    hdr_cells[1].text = 'Name'
    hdr_cells[2].text = 'Correct Answers'
    hdr_cells[3].text = 'Incorrect Answers'
    hdr_cells[4].text = 'Correct Percentage'
    
    for result in results:
        row_cells = table.add_row().cells
        row_cells[0].text = str(result.id)
        row_cells[1].text = result.author.username
        row_cells[2].text = str(result.correct_answers)
        row_cells[3].text = str(result.incorrect_answers)
        row_cells[4].text = f"{result.correct_percentage:.2f}%"
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{quiz.name}_results.docx"'
    doc.save(response)
    return response




def register(request):
    if request.method == 'POST':
        first_name = request.POST.get('frist_name')
        last_name = request.POST.get('last_name')
        username = request.POST.get('username')
        password = request.POST.get('password')
        password_confirm = request.POST.get('password-confirm')

        if password == password_confirm:
            if User.objects.filter(username=username).exists():
                messages.error(request, 'The username already exists')
            else:
                user = User.objects.create_user(username=username, password=password)
                user.first_name = first_name
                user.last_name = last_name
                user.username = username
                user.save()

                login(request, user)
                return redirect('quizList') 
        #else:
        #    return h(request, 'Passwords do not match')

    return render(request, 'register.html')
    


def login_view(request):
    if request.method == "POST":
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(username=username, password=password)
            if user is not None:
                login(request, user)
                return redirect(reverse('quizList'))
            else:
                # Handle invalid login
                pass
    else:
        form = AuthenticationForm()
    return render(request, 'login.html', {'form': form})
    
@login_required(login_url='login')
def logout_view(request):
    logout(request)
    return redirect('login')
    

def my_view(request):
    context = {
        'user': request.user
    }
    return render(request, 'base.html', context)


def error_handler(request, exception=None):
    
    
    error_list = ['error1.html', 'error2.html', 'error3.html']
    current_error = error_list.choice(error_list)
    
    
    return render(request, current_error, status=exception.status_code if exception else 500)

